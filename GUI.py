from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Text Scanner Output', 0)

p = document.add_paragraph('A plain paragraph having some ')
document.save('output.docx')